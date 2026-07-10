"""Integration tests for the document parsers (#13): ``extract_document`` per file type.

The per-type parsers (PDF/Word/Excel/text) preserve text + location so a later step can cite where
content came from. Folder → manifest ingestion (structure tree + located text) is covered by
``test_incremental_ingestion``; the corpus is vectorless, so there is no chunking seam.

Fixtures are synthesized with the real writer libraries (python-docx / openpyxl / reportlab) so the
parsers run against genuine files, not hand-rolled bytes.
"""

import pytest
from httpx import AsyncClient

pytestmark = pytest.mark.asyncio


# --------------------------- fixture builders ---------------------------

def _make_docx(path):
    from docx import Document

    d = Document()
    d.add_heading("Contrato", level=1)
    d.add_paragraph("Parte A contrata Parte B pelo prazo de 12 meses.")
    d.core_properties.author = "Ana Jurídica"
    d.save(str(path))


def _make_xlsx(path):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Vendas"
    ws.append(["mes", "receita"])
    ws.append(["jan", 1000])
    wb.save(str(path))


def _make_pdf(path):
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(str(path))
    c.drawString(72, 720, "Relatorio de vendas janeiro")
    c.save()


# --------------------------- Seam 1: parsers ---------------------------

class TestParsers:
    def test_pdf_text_extracted(self, tmp_path):
        from src.app.core.ingestion.parsers import extract_document

        p = tmp_path / "rel.pdf"
        _make_pdf(p)
        doc = extract_document(str(p))
        assert doc.doc_type == "pdf"
        assert "vendas" in " ".join(s.text for s in doc.sections).lower()

    def test_docx_sections_author_and_headings(self, tmp_path):
        from src.app.core.ingestion.parsers import extract_document

        p = tmp_path / "c.docx"
        _make_docx(p)
        doc = extract_document(str(p))
        assert doc.doc_type == "docx"
        assert doc.author == "Ana Jurídica"
        joined = " ".join(s.text for s in doc.sections)
        assert "Parte A" in joined
        assert any(s.location == "Contrato" for s in doc.sections)  # heading became a location

    def test_xlsx_sheet_preserves_schema(self, tmp_path):
        from src.app.core.ingestion.parsers import extract_document

        p = tmp_path / "v.xlsx"
        _make_xlsx(p)
        doc = extract_document(str(p))
        assert doc.doc_type == "xlsx"
        section = doc.sections[0]
        assert section.location == "planilha Vendas"
        assert "mes" in section.text and "jan" in section.text  # header + row preserved

    def test_unsupported_extension_raises(self, tmp_path):
        from src.app.core.ingestion.parsers import UnsupportedDocumentError, extract_document

        p = tmp_path / "x.zip"
        p.write_bytes(b"PK\x03\x04")
        with pytest.raises(UnsupportedDocumentError):
            extract_document(str(p))


