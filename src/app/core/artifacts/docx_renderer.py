"""Render an ArtifactSpec to a Word (.docx) report, applying a Template's visual identity."""

from typing import Optional

from docx import Document
from docx.shared import RGBColor

from src.app.core.artifacts.spec import ArtifactSpec, Template, claim_suffix


def render_docx(spec: ArtifactSpec, path: str, template: Optional[Template] = None) -> str:
    """Write ``spec`` as a .docx report at ``path`` and return the path.

    Each claim is rendered with its traceability suffix — the source, or a visible unsourced
    marker. Heading colour/fonts come from the template so content stays presentation-agnostic.
    """
    template = template or Template()
    doc = Document()

    title = doc.add_heading(spec.title, level=0)
    for run in title.runs:
        run.font.name = template.heading_font
        run.font.color.rgb = RGBColor.from_string(template.primary_color)
    if spec.subtitle:
        doc.add_paragraph(spec.subtitle)

    for section in spec.sections:
        heading = doc.add_heading(section.heading, level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor.from_string(template.primary_color)
        for claim in section.claims:
            paragraph = doc.add_paragraph(claim.text)
            marker = paragraph.add_run(claim_suffix(claim))
            marker.italic = claim.source is not None
            marker.bold = claim.source is None  # unsourced stands out

    doc.save(path)
    return path
