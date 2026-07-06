"""Document parsers: extract structured text from PDF, Word, Excel, and plain-text files.

Each parser preserves *location* so a chunk can later cite where it came from (page, sheet,
heading). PDF is OCR-ready: a page that yields no extractable text is flagged ``needs_ocr`` so a
scanned document is surfaced rather than silently dropped — actual OCR is an opt-in path
(``settings.OCR_ENABLED`` + tesseract), kept out of the default install so CI stays hermetic.
"""

import csv
import io
import os
from typing import List, Optional

from pydantic import BaseModel, Field

_TEXT_EXTS = {".txt", ".md", ".csv", ".log", ".json"}


class ParsedSection(BaseModel):
    """A located slice of a document (one page, one sheet, one heading-delimited block)."""

    location: str
    text: str
    needs_ocr: bool = False


class ParsedDocument(BaseModel):
    """A parsed document: its type, optional author, and located sections."""

    path: str
    doc_type: str  # pdf | docx | xlsx | text
    author: Optional[str] = None
    sections: List[ParsedSection] = Field(default_factory=list)


class UnsupportedDocumentError(ValueError):
    """Raised when a file's extension has no registered parser."""


def extract_document(path: str) -> ParsedDocument:
    """Parse a document into located sections, dispatching by file extension."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _parse_pdf(path)
    if ext == ".docx":
        return _parse_docx(path)
    if ext == ".xlsx":
        return _parse_xlsx(path)
    if ext in _TEXT_EXTS:
        return _parse_text(path)
    raise UnsupportedDocumentError(f"Sem parser para a extensão '{ext}'.")


def is_supported(path: str) -> bool:
    """Whether ``path`` has a registered parser."""
    ext = os.path.splitext(path)[1].lower()
    return ext in {".pdf", ".docx", ".xlsx", *_TEXT_EXTS}


def _parse_text(path: str) -> ParsedDocument:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    return ParsedDocument(path=path, doc_type="text", sections=[ParsedSection(location="conteúdo", text=text)])


def _parse_pdf(path: str) -> ParsedDocument:
    from pypdf import PdfReader

    reader = PdfReader(path)
    author = None
    try:
        author = reader.metadata.author if reader.metadata else None
    except Exception:  # noqa: BLE001 - metadata is best-effort
        author = None

    sections: List[ParsedSection] = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            text = (page.extract_text() or "").strip()
        except Exception:  # noqa: BLE001 - a malformed page must not abort the whole document
            text = ""
        # A page with no extractable text is (almost certainly) scanned — flag for OCR.
        sections.append(ParsedSection(location=f"página {i}", text=text, needs_ocr=not bool(text)))
    return ParsedDocument(path=path, doc_type="pdf", author=author, sections=sections)


def _parse_docx(path: str) -> ParsedDocument:
    from docx import Document

    doc = Document(path)
    author = None
    try:
        author = doc.core_properties.author or None
    except Exception:  # noqa: BLE001
        author = None

    sections: List[ParsedSection] = []
    current_location = "corpo"
    buffer: List[str] = []

    def _flush() -> None:
        if buffer:
            sections.append(ParsedSection(location=current_location, text="\n".join(buffer).strip()))
            buffer.clear()

    for para in doc.paragraphs:
        style = (para.style.name if para.style else "") or ""
        if style.startswith("Heading") and para.text.strip():
            _flush()
            current_location = para.text.strip()
        elif para.text.strip():
            buffer.append(para.text.strip())
    _flush()

    # Tables preserved as pipe-delimited rows so structure survives.
    for t_idx, table in enumerate(doc.tables, start=1):
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        if rows:
            sections.append(ParsedSection(location=f"tabela {t_idx}", text="\n".join(rows)))

    return ParsedDocument(path=path, doc_type="docx", author=author, sections=sections)


def _parse_xlsx(path: str, max_rows: int = 2000) -> ParsedDocument:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    sections: List[ParsedSection] = []
    for ws in wb.worksheets:
        out = io.StringIO()
        writer = csv.writer(out)
        for r_idx, row in enumerate(ws.iter_rows(values_only=True)):
            if r_idx >= max_rows:
                break
            writer.writerow(["" if v is None else v for v in row])
        text = out.getvalue().strip()
        if text:
            sections.append(ParsedSection(location=f"planilha {ws.title}", text=text))
    wb.close()
    return ParsedDocument(path=path, doc_type="xlsx", sections=sections)
